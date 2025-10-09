from flask import Flask, render_template, request, render_template_string, send_file, flash, redirect, url_for, jsonify
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import tempfile
from werkzeug.utils import secure_filename
from datetime import datetime
import uuid
import base64

app = Flask(__name__)
app.secret_key = 'your-secret-key-here'
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max file size
app.config['UPLOAD_FOLDER'] = 'temp_uploads'

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def set_landscape_orientation(doc):
    """Set document to landscape orientation"""
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(42.0)   # panjang A3
        section.page_height = Cm(29.7)  # lebar A3
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

def add_page_break(doc):
    """Add page break to document"""
    paragraph = doc.add_paragraph()
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

def create_title_header(doc, title):
    """Create title header matching the template"""
    title_para = doc.add_paragraph(title)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.runs[0]
    run.font.size = Pt(16)
    run.bold = True
    p_format = title_para.paragraph_format
    p_format.space_before = Pt(0)   # jarak sebelum judul
    p_format.space_after = Pt(24) 
    return title_para

def add_image_to_cell(cell, image_path, width=None):
    """Add image to table cell with proper alignment and error handling"""
    if image_path and os.path.exists(image_path):
        # Clear existing content
        cell.text = ''
        paragraph = cell.paragraphs[0]
        paragraph.clear()
        
        run = paragraph.add_run()
        try:
            picture = run.add_picture(image_path)  # tambahkan dulu tanpa ukuran
            
            picture.width = Cm(5.2)
            picture.height = Cm(7.1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_format = paragraph.paragraph_format
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(8)

        except Exception as e:
            # If image fails to load, add error text
            cell.text = f"Image Error: {str(e)[:100]}"
            print(f"Error adding image {image_path}: {str(e)}")
    else:
        cell.text = "Image not found"
        

def add_label_box(cell, text, width_cm=5, height_cm=0.6):
    # buat tabel 1x1
    table = cell.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # atur ukuran
    table.columns[0].width = Cm(width_cm)
    table.rows[0].height = Cm(height_cm)
    
    # isi teks
    inner_cell = table.cell(0, 0)
    inner_cell.width = Cm(width_cm)
    inner_cell.height = Cm(height_cm)

    paragraph = inner_cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)

    return table

def set_cell_margin(cell, top=0, start=0, bottom=0, end=0):
    """Atur margin cell (dalam twips, 1 pt ≈ 20 twips)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = OxmlElement(f"w:{m}")
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def remove_table_borders(table):
        tbl = table._element
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')   # nil = tidak ada border
            tblBorders.append(border)
        tbl.tblPr.append(tblBorders)

def generate_word_document(form_data, uploaded_files):
    """Generate Word document in landscape orientation with improved structure"""
    doc = Document()
    
    # Set landscape orientation and margins
    set_landscape_orientation(doc)
    
    # PAGE 1 - GPON & FTM DATA (Landscape Layout)
    section = doc.sections[0] 
    header = section.header

    # kosongkan isi header dulu
    header.is_linked_to_previous = False
    header.paragraphs[0].clear()

    # tambahkan judul ke header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    run = paragraph.add_run(form_data.get('judul_laporan'))

    run.bold = True
    run.font.size = Pt(16)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    

    # Create main table for both GPON and FTM side by side (better use of landscape space)
    main_table = doc.add_table(rows=1, cols=2)
    main_table.autofit = False
    remove_table_borders(main_table)

    for row in main_table.rows:
        row.cells[0].width = Cm(18)
        row.cells[1].width = Cm(18)
    
    # GPON Section (Left side)
    gpon_cell = main_table.cell(0, 0)
    gpon_para = gpon_cell.paragraphs[0]
    gpon_para.clear()
    
    # GPON Info Box
    gpon_info_table = gpon_cell.add_table(rows=1, cols=1)
    gpon_info_table.style = 'Table Grid'
    gpon_info_cell = gpon_info_table.cell(0, 0)
    gpon_info_table.autofit =False
    gpon_info_table.columns[0].width = Cm(9)
    gpon_info_cell.width = Cm(8)
    gpon_info_table.alignment = WD_TABLE_ALIGNMENT.CENTER 

    gpon_info_para = gpon_info_cell.paragraphs[0]
    gpon_info_para.add_run(f"{form_data.get('sto_gpon',)}").bold = True
    gpon_info_para.add_run(f"\n{form_data.get('ruangan_gpon')}").bold = True
    gpon_info_para.add_run(f"\n{form_data.get('koordinat_gpon')}").bold = True
    gpon_info_para.add_run(f"\n{form_data.get('IP_Address')}").bold = True
    gpon_info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # GPON Images Grid (2x2) - optimized for landscape
    gpon_img_table = gpon_cell.add_table(rows=2, cols=2)
    gpon_img_table.style = 'Table Grid'
    gpon_img_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # pastikan tiap cell juga ikut
    for row in gpon_img_table.rows:
        row.cells[0].width = Cm(6.5)
        row.cells[1].width = Cm(6.5)
    
    # GPON Images with larger size for landscape
    gpon_images = [
        ('foto_gpon_1', 'lemari_gpon'),  # Remove tuple wrapper
        ('foto_gpon_2', 'keterangan_gpon_2'),
        ('foto_gpon_3', 'card_gpon'),
        ('foto_gpon_4', 'port_gpon')
    ]

    positions = [(0, 0), (0, 1), (1, 0), (1, 1)]
    for i, (img_key, text_key) in enumerate(gpon_images):
        row, col = positions[i]
        cell = gpon_img_table.cell(row, col)
        cell.text = ""  # Clear default content
        
        # Create nested table for image + caption layout
        nested = cell.add_table(rows=2, cols=1)
        nested.autofit = False
        nested.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Image row
        img_cell = nested.cell(0, 0)
        if uploaded_files.get(img_key):
            add_image_to_cell(img_cell, uploaded_files.get(img_key), width=Inches(2.0))
        
        # Caption row
        text_cell = nested.cell(1, 0)
        add_label_box(text_cell, form_data.get(text_key, ""))
        set_cell_margin(cell, bottom=50)

    # FTM Section (Right side)
    ftm_cell = main_table.cell(0, 1)
    ftm_para = ftm_cell.paragraphs[0]
    ftm_para.clear()
    
    # FTM Info Box
    ftm_info_table = ftm_cell.add_table(rows=1, cols=1)
    ftm_info_table.style = 'Table Grid'
    ftm_info_cell = ftm_info_table.cell(0, 0)
    ftm_info_table.autofit =False
    ftm_info_table.columns[0].width = Cm(9)
    ftm_info_cell.width = Cm(8)
    ftm_info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    ftm_info_para = ftm_info_cell.paragraphs[0]
    ftm_info_para.add_run(f"{form_data.get('sto_ftm', )}").bold = True
    ftm_info_para.add_run(f"\n{form_data.get('ruangan_ftm')}").bold = True
    ftm_info_para.add_run(f"\n{form_data.get('koordinat_ftm')}").bold = True
    ftm_info_para.add_run(f"\n{form_data.get('kode_ftm')}").bold = True
    ftm_info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # FTM Images Grid (2x2)
    ftm_img_table = ftm_cell.add_table(rows=2, cols=2)
    ftm_img_table.style = 'Table Grid'
    ftm_img_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # pastikan tiap cell juga ikut
    for row in ftm_img_table.rows:
        row.cells[0].width = Cm(6.5)
        row.cells[1].width = Cm(6.5)
    
    ftm_images = [
        ('foto_ftm_1', 'lemari_ftm'),  # Remove tuple wrapper
        ('foto_ftm_2', 'keterangan_ftm_2'),
        ('foto_ftm_3', 'otb_ftm'),
        ('foto_ftm_4', 'slot_port_ftm')
    ]

    positions = [(0, 0), (0, 1), (1, 0), (1, 1)]
    for i, (img_key, text_key) in enumerate(ftm_images):
        row, col = positions[i]
        cell = ftm_img_table.cell(row, col)
        cell.text = ""  # Clear default content
        
        # Create nested table for image + caption layout
        nested = cell.add_table(rows=2, cols=1)
        nested.autofit = False
        nested.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Image row
        img_cell = nested.cell(0, 0)
        if uploaded_files.get(img_key):
            add_image_to_cell(img_cell, uploaded_files.get(img_key), width=Inches(2.0))
        
        # Caption row
        text_cell = nested.cell(1, 0)
        add_label_box(text_cell, form_data.get(text_key, ""))
        
        set_cell_margin(cell, bottom=50)

    # PAGE 2 - FTM Detail & ODC (Landscape optimized)
    add_page_break(doc)
    
    # Create landscape layout for FTM Detail and ODC
    page2_table = doc.add_table(rows=1, cols=2)
    page2_table.autofit = False
    remove_table_borders(page2_table)

    for row in page2_table.rows:
        row.cells[0].width = Cm(18.5)
        row.cells[1].width = Cm(18.5)
    
    # FTM Detail Section (Left)
    ftm_detail_cell = page2_table.cell(0, 0)
    ftm_detail_para = ftm_detail_cell.paragraphs[0]
    ftm_detail_para.clear()
    
    # FTM detail Info Box
    ftm_detail_info_table = ftm_detail_cell.add_table(rows=1, cols=1)
    ftm_detail_info_table.style = 'Table Grid'
    ftm_detail_info_cell = ftm_detail_info_table.cell(0, 0)
    ftm_detail_info_table.autofit =False
    ftm_detail_info_table.columns[0].width = Cm(9)
    ftm_detail_info_cell.width = Cm(8)
    ftm_detail_info_table.alignment = WD_TABLE_ALIGNMENT.CENTER 

    # FTM Detail Info
    ftm_detail_info_para = ftm_detail_info_cell.paragraphs[0]
    ftm_detail_info_para.add_run(f" {form_data.get('ruang_ftm')}").bold = True
    ftm_detail_info_para.add_run(f"\n{form_data.get('koordinat_ftm_detail')}").bold = True
    ftm_detail_info_para.add_run(f"\n{form_data.get('kode_ftm_detail')}").bold = True
    ftm_detail_info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Bagian atas: 2 gambar
    ftm_detail_top_table = ftm_detail_cell.add_table(rows=1, cols=2)
    ftm_detail_top_table.style = 'Table Grid'
    ftm_detail_top_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in ftm_detail_top_table.rows:
        row.cells[0].width = Cm(6.5)
        row.cells[1].width = Cm(6.5)

    # Bagian bawah: 3 gambar
    ftm_detail_bottom_table = ftm_detail_cell.add_table(rows=1, cols=3)
    ftm_detail_bottom_table.style = 'Table Grid'
    for row in ftm_detail_bottom_table.rows:
        row.cells[0].width = Cm(5.9)
        row.cells[1].width = Cm(5.9)
        row.cells[2].width = Cm(5.9)

    
    # === Daftar gambar ===
    ftm_detail_images = [
        ('foto_ftm_detail_1', 'no_lemari_ftm'),
        ('foto_ftm_detail_2', 'keterangan_ftm_detail_2'),
        ('foto_ftm_detail_3', 'no_otb'),
        ('foto_ftm_detail_4', 'no_panel'),
        ('foto_ftm_detail_5', 'no_port_core')
    ]
    
    for i in range(2):
        img_key, text_key = ftm_detail_images[i]
        cell = ftm_detail_top_table.cell(0, i)
        cell.text = ""  # kosongkan biar gak ada paragraf default

        # bikin nested table di dalam cell utama
        nested = cell.add_table(rows=2, cols=1)
        nested.autofit = False
        nested.alignment = WD_TABLE_ALIGNMENT.CENTER

        # baris gambar
        img_cell = nested.cell(0, 0)
        if uploaded_files.get(img_key):
            add_image_to_cell(img_cell, uploaded_files.get(img_key))

        # baris teks
        text_cell = nested.cell(1, 0)
        add_label_box(text_cell, form_data.get(text_key))
        set_cell_margin(cell, bottom=50)
        
        # === Masukkan gambar ke tabel bawah (3 gambar) ===
    for i in range(3):
        img_key, text_key = ftm_detail_images[i+2]
        cell = ftm_detail_bottom_table.cell(0, i)
        cell.text = ""  # kosongkan biar gak ada paragraf default

        # bikin nested table di dalam cell utama
        nested = cell.add_table(rows=2, cols=1)
        nested.autofit = False
        nested.alignment = WD_TABLE_ALIGNMENT.CENTER

        # baris gambar
        img_cell = nested.cell(0, 0)
        if uploaded_files.get(img_key):
            add_image_to_cell(img_cell, uploaded_files.get(img_key))

        # baris teks
        text_cell = nested.cell(1, 0)
        add_label_box(text_cell, form_data.get(text_key))
        set_cell_margin(cell, bottom=50)
    
    # ODC Section (Right)
    odc_cell = page2_table.cell(0, 1)
    odc_para = odc_cell.paragraphs[0]
    odc_para.clear()


    # ODC Info
    odc_info_table =  odc_cell.add_table(rows=1, cols=1)
    odc_info_table.style = 'Table Grid'
    odc_info_cell = odc_info_table.cell(0, 0)
    odc_info_table.autofit =False
    odc_info_table.columns[0].width = Cm(9)
    odc_info_cell.width = Cm(8)
    odc_info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    odc_info_para = odc_info_cell.paragraphs[0]
    odc_info_para.add_run(f"{form_data.get('odc_nama')}").bold = True
    odc_info_para.add_run(f"\n{form_data.get('kode_odc')}").bold = True
    odc_info_para.add_run(f"\n{form_data.get('koordinat_odc')}").bold = True
    odc_info_para.add_run(f"\n{form_data.get('lokasi_odc')}").bold = True
    odc_info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Bagian atas: 2 gambar
    odc_top_table = odc_cell.add_table(rows=1, cols=2)
    odc_top_table.style = 'Table Grid'
    odc_top_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in odc_top_table.rows:
        row.cells[0].width = Cm(6)
        row.cells[1].width = Cm(6.5)

    # Bagian bawah: 3 gambar
    odc_bottom_table = odc_cell.add_table(rows=1, cols=3)
    odc_bottom_table.style = 'Table Grid'
    for row in odc_bottom_table.rows:
        row.cells[0].width = Cm(5.9)
        row.cells[1].width = Cm(5.9)
        row.cells[2].width = Cm(5.9)

    odc_images = [
        ('foto_odc_1', ('keterangan_odc_1',)),
        ('foto_odc_2', ('keterangan_odc_2',)),
        ('foto_odc_3', ('no_in_tray', 'no_tray')),
        ('foto_odc_4', ('no_port_core_odc', 'no_core_port_odc')),
        ('foto_odc_5', ('hasil_ukur', 'feeder'))
    ]
    
    # === Masukkan gambar ke tabel atas (2 gambar) ===
    for i in range(2):
        img_key, text_keys = odc_images[i]
        cell = odc_top_table.cell(0, i)
        cell.text = ""

        nested = cell.add_table(rows=2, cols=1)  # harus 2 baris (gambar + teks)
        nested.autofit = False
        nested.alignment = WD_TABLE_ALIGNMENT.CENTER

        img_cell = nested.cell(0, 0)
        if uploaded_files.get(img_key):
            add_image_to_cell(img_cell, uploaded_files.get(img_key))

        text_cell = nested.cell(1, 0)
        for text_key in text_keys:   # loop karena text_keys tuple
            add_label_box(text_cell, form_data.get(text_key, ""))
            set_cell_margin(cell, bottom=50)


    # === Masukkan gambar ke tabel bawah (3 gambar) ===
    for i in range(3):
        img_key, text_keys = odc_images[i+2]
        cell = odc_bottom_table.cell(0, i)
        cell.text = ""

        nested = cell.add_table(rows=2, cols=1)  # sama, 2 baris
        nested.autofit = False
        nested.alignment = WD_TABLE_ALIGNMENT.CENTER

        img_cell = nested.cell(0, 0)
        if uploaded_files.get(img_key):
            add_image_to_cell(img_cell, uploaded_files.get(img_key))

        text_cell = nested.cell(1, 0)
        for text_key in text_keys:
            add_label_box(text_cell, form_data.get(text_key, ""))
            set_cell_margin(cell, bottom=50)



    ####################################
    
    # PAGE 3 - SPL & ODP (Landscape optimized)
    add_page_break(doc)
    
    page3_table = doc.add_table(rows=1, cols=2)
    page3_table.autofit = False
    remove_table_borders(page3_table)

    for row in page3_table.rows:
        row.cells[0].width = Cm(18.5)
        row.cells[1].width = Cm(18.5)

    left_cell = page3_table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_para.clear()

    # Info ODC
    odc_info_table = left_cell.add_table(rows=1, cols=1)
    odc_info_table.style = 'Table Grid'
    odc_cell_inner = odc_info_table.cell(0, 0)
    odc_info_table.autofit = False
    odc_info_table.columns[0].width = Cm(9)
    odc_cell_inner.width = Cm(8)
    odc_info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    odc_para = odc_cell_inner.paragraphs[0]
    odc_para.add_run(f"{form_data.get('odc_hal3')}").bold = True
    odc_para.add_run(f"\n{form_data.get('detail_odc_hal3')}").bold = True
    odc_para.add_run(f"\n{form_data.get('koordinat_odc_hal3')}").bold = True
    odc_para.add_run(f"\n{form_data.get('lokasi_odc_hal3')}").bold = True
    odc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Foto SPL/ODC (2x2 grid)
    spl_img_table = left_cell.add_table(rows=2, cols=2)
    spl_img_table.style = 'Table Grid'
    spl_img_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in spl_img_table.rows:
        row.cells[0].width = Cm(6.5)
        row.cells[1].width = Cm(6.5)

    spl_images = [
        ('foto_spl_1', ('keterangan_spl_1', 'spl_1')),
        ('foto_spl_2', ('ukur_spl','hasil_ukur_spl')),
        ('foto_out', ('in_out', 'keterangan_out')),
        ('foto_port', ('keterangan_port','keterangan_core'))
    ]
    for i, (img_key, text_keys) in enumerate(spl_images):
        row, col = positions[i]
        cell = spl_img_table.cell(row, col)
        cell.text = ""
        
        # Buat nested table
        nested = cell.add_table(rows=2, cols=1)
        nested.autofit = False
        nested.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Gambar
        img_cell = nested.cell(0, 0)
        if uploaded_files.get(img_key):
            add_image_to_cell(img_cell, uploaded_files.get(img_key))
        
        # Text labels
        text_cell = nested.cell(1, 0)
        for text_key in text_keys:
            add_label_box(text_cell, form_data.get(text_key, ""))
        
        set_cell_margin(cell, bottom=50)


    # ---------------- KANAN : ODP ---------------- 
    right_cell = page3_table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_para.clear()

    # Info ODP
    odp_info_table = right_cell.add_table(rows=1, cols=1)
    odp_info_table.style = 'Table Grid'
    odp_cell_inner = odp_info_table.cell(0, 0)
    odp_info_table.autofit = False
    odp_info_table.columns[0].width = Cm(9)
    odp_cell_inner.width = Cm(8)
    odp_info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    odp_para = odp_cell_inner.paragraphs[0]
    odp_para.add_run(f"{form_data.get('odp_hal3')} - ").bold = True
    odp_para.add_run(f"\n{form_data.get('detail_odp_hal3')}").bold = True
    odp_para.add_run(f"\n{form_data.get('lokasi_odp_hal3')}").bold = True
    odp_para.add_run(f"\n{form_data.get('koordinat_odp_hal3')}").bold = True
    odp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Foto ODP & lainnya (3x2 grid)
    odp_img_table = right_cell.add_table(rows=2, cols=3)
    odp_img_table.style = 'Table Grid'
    odp_img_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in odp_img_table.rows:
        row.cells[0].width = Cm(6)
        row.cells[1].width = Cm(6)
        row.cells[2].width = Cm(6)

    odp_images = [
        ('foto_odp_1', ('keterangan_odp_1', 'keterangan_odp_12')),
        ('foto_odp_2', ('keterangan_odp_2', 'keterangan_odp_21')),
        ('foto_spl_3', ('keterangan_spl_3', 'keterangan_spl_31')),
        ('foto_qr_1', ('keterangan_qr_1', 'keterangan_qr_12')),
        ('foto_qr_2', ('keterangan_qr_2', 'keterangan_qr_21')),
        ('foto_valin', ('keterangan_valin', 'keterangan_valin1'))
    ]

    positions = [(0,0), (0,1), (0,2), (1,0), (1,1), (1,2)]
    for i, (img_key, text_keys) in enumerate(odp_images):
        row, col = positions[i]
        cell = odp_img_table.cell(row, col)
        cell.text = ""
        
        # Buat nested table
        nested = cell.add_table(rows=2, cols=1)
        nested.autofit = False
        nested.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Gambar
        img_cell = nested.cell(0, 0)
        if uploaded_files.get(img_key):
            add_image_to_cell(img_cell, uploaded_files.get(img_key))
        
        # Text labels
        text_cell = nested.cell(1, 0)
        for text_key in text_keys:
            add_label_box(text_cell, form_data.get(text_key, ""))
        
        set_cell_margin(cell, bottom=50)

    # PAGE 4 - 8 PORT (Landscape optimized - 4x2 layout)
    add_page_break(doc)

    page4_table = doc.add_table(rows=2, cols=1)
    page4_table.autofit = False
    remove_table_borders(page4_table)

    odp4_cell = page4_table.cell(0,0)
    odp4_para = odp4_cell.add_paragraph()
    odp4_para.clear()

    odp4_info_table = odp4_cell.add_table(rows=1, cols=1)
    odp4_info_table.style = 'Table Grid'
    odp4_info_cell = odp4_info_table.cell(0, 0)
    odp4_info_table.autofit = False
    odp4_info_table.columns[0].width = Cm(9)
    odp4_info_cell.width = Cm(8)
    odp4_info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # ODP Info for page 4
    odp4_info = odp4_info_cell.paragraphs[0]
    odp4_info.add_run(f"{form_data.get('odp_hal4','')}").bold = True
    odp4_info.add_run(f"\n{form_data.get('odp1_hal4','')}").bold = True
    odp4_info.add_run(f"\n{form_data.get('lokasi_odp_hal4','')}").bold = True
    odp4_info.add_run(f"\n{form_data.get('detail_odp_hal4','')}").bold = True
    odp4_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    port_table = odp4_cell.add_table(rows=2, cols=4)
    port_table.style = 'Table Grid'
    port_table.autofit = False
    for row in port_table.rows:
        for cell in row.cells:
            cell.width = Cm(8)
    port_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # All 8 port images
    port_images = [
        ('foto_port_1', 'keterangan_port_1'),
        ('foto_port_2', 'keterangan_port_2'),
        ('foto_port_3', 'keterangan_port_3'),
        ('foto_port_4', 'keterangan_port_4'),
        ('foto_port_5', 'keterangan_port_5'),
        ('foto_port_6', 'keterangan_port_6'),
        ('foto_port_7', 'keterangan_port_7'),
        ('foto_port_8', 'keterangan_port_8')
    ]
    
    positions = [(0,0), (0,1), (0,2), (0,3), (1,0), (1,1), (1,2), (1,3)]
    for i, (img_key, text_key) in enumerate(port_images):
        row, col = positions[i]
        cell = port_table.cell(row, col)

        # Tambah gambar jika ada
        if uploaded_files.get(img_key):
            add_image_to_cell(cell, uploaded_files.get(img_key), width=Inches(2.0))

        # Tambah caption teks + label port
        caption = form_data.get(text_key, '')
        add_label_box(cell, caption)

        # Tambah margin bawah
        set_cell_margin(cell, bottom=50)

    
    # PAGE 5 - DOCUMENTATION (Landscape optimized)
    add_page_break(doc)
    pa = doc.add_paragraph()
    pa.paragraph_format.space_after = Cm (0.5)
    
    doc_images = [
        'foto_dok_1',
        'foto_dok_2',
        'foto_dok_3'
    ]
    
    for img_key in doc_images:
        if uploaded_files.get(img_key):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(uploaded_files[img_key], height=Inches(2.7))
    
    # PAGE 6 - DENAH LOKASI (Landscape optimized)
    add_page_break(doc)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Cm(1.5)
    
    # Denah image - larger for landscape
    if uploaded_files.get('foto_denah'):
        denah_para = doc.add_paragraph()
        denah_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        denah_run = denah_para.add_run()
        try:
            denah_run.add_picture(uploaded_files.get('foto_denah'), width=Inches(13))  # Larger for landscape
        except Exception as e:
            denah_para.text = f"Error loading denah image: {str(e)}"
    else:
        denah_para = doc.add_paragraph("Denah image not available")
        denah_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

@app.route('/')
def index():
    """Main page"""
    return render_template("index_v7.html")


@app.route('/generate', methods=['POST'])
def generate_report():
    """Generate Word document from form submission"""
    try:
        form_data = request.form.to_dict()
        uploaded_files = {}
        temp_files = []
        
        # Process uploaded files
        for field_name in request.files:
            file = request.files[field_name]
            if file and file.filename != '' and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                unique_filename = f"{uuid.uuid4().hex}_{filename}"
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
                file.save(filepath)
                uploaded_files[field_name] = filepath
                temp_files.append(filepath)
        
        if len(uploaded_files) == 0:
            flash('Harap upload minimal satu foto!', 'error')
            return redirect(url_for('index'))
        
        # Generate Word document in landscape
        doc = generate_word_document(form_data, uploaded_files)
        
        # Save document to temporary file
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_filename = f"Laporan_Evidence_Landscape_{form_data.get('judul_laporan', 'Report')}_{timestamp}.docx"
        
        # Create temporary file for the document
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            doc.save(tmp_file.name)
            
            # Clean up uploaded files
            for temp_file in temp_files:
                try:
                    os.remove(temp_file)
                except:
                    pass
            
            # Return the file
            return send_file(
                tmp_file.name,
                as_attachment=True,
                download_name=report_filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
    
    except Exception as e:
        # Clean up uploaded files in case of error
        for temp_file in temp_files:
            try:
                os.remove(temp_file)
            except:
                pass
        
        flash(f'Terjadi kesalahan: {str(e)}', 'error')
        return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
                